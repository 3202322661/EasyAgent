"""
Word 文档操作工具模块。

提供完整的 .docx 文档创建、编辑和读取能力：
  读取: read_word_text, read_word_tables, read_word_info, list_doc_styles
  创建: create_word_document, add_paragraph, add_heading, add_table, add_page_break
  格式: set_paragraph_format, set_run_format, set_page_margins
"""

import logging
import os
import shutil
import tempfile
from typing import Any, List, Optional, Union

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from tools._utils import parse_color as _parse_rgb, parse_alignment

logger = logging.getLogger(__name__)

# ────────────────────────── 常量映射 ──────────────────────────

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

TABLE_ALIGN_MAP = {
    "left":   WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right":  WD_TABLE_ALIGNMENT.RIGHT,
}


# ────────────────────────── 颜色解析（python-docx 专用） ──────────────────────────

def _parse_docx_color(color_val: Optional[str]) -> Optional[RGBColor]:
    """将颜色字符串解析为 python-docx 的 RGBColor 对象。"""
    rgb = _parse_rgb(color_val)
    if rgb is None:
        return None
    return RGBColor(*rgb)


# ────────────────────────── 内部辅助 ──────────────────────────

def _safe_open_doc(file_path: str) -> Document:
    """安全地打开或创建 Word 文档。"""
    if not os.path.exists(file_path):
        return Document()
    try:
        return Document(file_path)
    except Exception as e:
        logger.warning("打开文档 '%s' 失败 (%s)，创建新文档。", file_path, e)
        return Document()


def _safe_save_doc(doc: Document, file_path: str) -> bool:
    """防御性保存：先写临时文件再替换，避免保存中断损坏原文件。"""
    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(tmp_path)
        shutil.copy2(tmp_path, file_path)
        return True
    except Exception as e:
        logger.error("保存 Word 文档 '%s' 失败: %s", file_path, e)
        return False
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def _set_run_font(run: Any, font_name: Optional[str] = None,
                  font_size: Optional[int] = None,
                  bold: Optional[bool] = None,
                  italic: Optional[bool] = None,
                  underline: Optional[bool] = None,
                  color: Optional[str] = None) -> None:
    """安全地设置 Run 对象的字体格式。"""
    if font_name:
        try:
            run.font.name = font_name
            # 同步设置中文字体
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass

    if font_size is not None:
        try:
            run.font.size = Pt(font_size)
        except Exception:
            pass
    if bold is not None:
        try:
            run.bold = bold
        except Exception:
            pass
    if italic is not None:
        try:
            run.italic = italic
        except Exception:
            pass
    if underline is not None:
        try:
            run.underline = underline
        except Exception:
            pass
    if color:
        try:
            parsed = _parse_docx_color(color)
            if parsed:
                run.font.color.rgb = parsed
        except Exception:
            pass


def _set_para_format(paragraph: Any, alignment: Optional[str] = None,
                     line_spacing: Optional[float] = None,
                     space_before: Optional[int] = None,
                     space_after: Optional[int] = None,
                     first_line_indent: Optional[int] = None,
                     left_indent: Optional[int] = None,
                     right_indent: Optional[int] = None) -> None:
    """安全地设置段落排版格式。"""
    pf = paragraph.paragraph_format
    if alignment:
        val = parse_alignment(alignment, ALIGN_MAP)
        if val is not None:
            pf.alignment = val
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if right_indent is not None:
        pf.right_indent = Pt(right_indent)


# ────────────────────────── 读取接口 ──────────────────────────

def read_word_text(file_path: str) -> str:
    """读取 Word 文档的纯文本内容，带段落编号与样式名称。

    Args:
        file_path: .docx 文件路径

    Returns:
        带编号的段落文本
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        lines: List[str] = []
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            lines.append(f"{i:4d} [{style_name}] {text}")

        non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
        result = "\n".join(lines)
        return f"文件 '{file_path}' 文本内容（{len(lines)} 段，{non_empty} 段非空）:\n{result}"

    except Exception as e:
        logger.exception("读取 Word 文本失败")
        return f"[ERROR] 读取 Word 文档失败: {e}"


def read_word_tables(file_path: str) -> str:
    """读取 Word 文档中所有表格的内容。

    Args:
        file_path: .docx 文件路径

    Returns:
        格式化的表格数据
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        if not doc.tables:
            return f"文件 '{file_path}' 中未找到表格。"

        parts: List[str] = []
        for t_idx, table in enumerate(doc.tables, 1):
            parts.append(f"--- 表格 {t_idx} ({len(table.rows)} 行 x {len(table.columns)} 列) ---")
            for r_idx, row in enumerate(table.rows, 1):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                parts.append(f"  行{r_idx}: {' | '.join(cells)}")
            parts.append("")

        return f"文件 '{file_path}' 的表格:\n" + "\n".join(parts)

    except Exception as e:
        logger.exception("读取 Word 表格失败")
        return f"[ERROR] 读取表格失败: {e}"


def read_word_info(file_path: str) -> str:
    """读取 Word 文档概要信息（段落数、表格数、样式等）。

    Args:
        file_path: .docx 文件路径

    Returns:
        文档概要
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        para_count = len(doc.paragraphs)
        non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
        table_count = len(doc.tables)

        styles_used: set[str] = set()
        for p in doc.paragraphs:
            if p.style:
                styles_used.add(p.style.name)

        lines = [
            f"文档: {file_path}",
            f"总段落数: {para_count}",
            f"非空段落: {non_empty}",
            f"表格数量: {table_count}",
            f"使用样式 ({len(styles_used)}): {', '.join(sorted(styles_used))}",
        ]
        return "\n".join(lines)

    except Exception as e:
        logger.exception("读取 Word 信息失败")
        return f"[ERROR] 读取文档信息失败: {e}"


def list_doc_styles(file_path: str) -> str:
    """列出 Word 文档中所有可用样式及类型。

    Args:
        file_path: .docx 文件路径

    Returns:
        样式列表
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        info: List[str] = []
        for style in doc.styles:
            if style.type is not None:
                info.append(f"  [{style.type.name}] {style.name}")
        return f"文档 '{file_path}' 的可用样式（{len(info)} 个）:\n" + "\n".join(info)

    except Exception as e:
        logger.exception("获取样式列表失败")
        return f"[ERROR] 获取样式列表失败: {e}"


# ────────────────────────── 创建/编辑接口 ──────────────────────────

def create_word_document(file_path: str, title: Optional[str] = None) -> str:
    """创建新的空白 Word 文档。

    Args:
        file_path: 保存路径 (.docx)
        title: 可选的居中标题

    Returns:
        操作结果
    """
    try:
        doc = Document()
        if title:
            para = doc.add_heading(title, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(file_path)
        return f"[SUCCESS] Word 文档已创建: '{file_path}'"
    except Exception as e:
        logger.exception("创建 Word 文档失败")
        return f"[ERROR] 创建文档失败: {e}"


def add_paragraph(file_path: str, text: str,
                  style: Optional[str] = None,
                  font_name: Optional[str] = None,
                  font_size: Optional[int] = None,
                  bold: Optional[bool] = None,
                  italic: Optional[bool] = None,
                  underline: Optional[bool] = None,
                  color: Optional[str] = None,
                  alignment: Optional[str] = None,
                  line_spacing: Optional[float] = None,
                  space_before: Optional[int] = None,
                  space_after: Optional[int] = None,
                  first_line_indent: Optional[int] = None) -> str:
    """向 Word 文档添加段落，支持完整的字体和排版设置。

    Args:
        file_path: 文档路径（不存在则自动创建）
        text: 段落文本
        style: 样式名，如 'Normal', 'Heading 1'
        font_name: 字体名称
        font_size: 字号（磅）
        bold: 是否加粗
        italic: 是否斜体
        underline: 是否下划线
        color: 文字颜色（名称或 #HEX）
        alignment: 对齐方式 (left/center/right/justify)
        line_spacing: 行距倍数
        space_before: 段前距（磅）
        space_after: 段后距（磅）
        first_line_indent: 首行缩进（磅）

    Returns:
        操作结果
    """
    try:
        doc = _safe_open_doc(file_path)
        para = doc.add_paragraph()
        run = para.add_run(text)

        _set_run_font(run, font_name=font_name, font_size=font_size,
                      bold=bold, italic=italic, underline=underline, color=color)
        _set_para_format(para, alignment=alignment, line_spacing=line_spacing,
                         space_before=space_before, space_after=space_after,
                         first_line_indent=first_line_indent)

        if style:
            try:
                para.style = doc.styles[style]
            except KeyError:
                pass  # 样式不存在时静默忽略

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 段落已添加至 '{file_path}'"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("添加段落失败")
        return f"[ERROR] 添加段落失败: {e}"


def add_heading(file_path: str, text: str, level: int = 1,
                font_name: Optional[str] = None,
                font_size: Optional[int] = None,
                color: Optional[str] = None,
                alignment: Optional[str] = None) -> str:
    """向 Word 文档添加标题（1-9 级）。

    Args:
        file_path: 文档路径
        text: 标题文字
        level: 标题级别 (1-9)
        font_name: 字体名称
        font_size: 字号（磅）
        color: 文字颜色
        alignment: 对齐方式

    Returns:
        操作结果
    """
    if not 1 <= level <= 9:
        return f"[ERROR] 标题级别必须为 1-9，收到: {level}"

    try:
        doc = _safe_open_doc(file_path)
        para = doc.add_paragraph()
        run = para.add_run(text)

        # 应用 Heading 样式
        style_name = f"Heading {level}"
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass

        if font_name or font_size or color:
            _set_run_font(run, font_name=font_name, font_size=font_size, color=color)
        if alignment:
            _set_para_format(para, alignment=alignment)

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 标题 (H{level}) 已添加至 '{file_path}'"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("添加标题失败")
        return f"[ERROR] 添加标题失败: {e}"


def add_table(file_path: str, data: list,
              headers: Optional[list] = None,
              font_name: Optional[str] = None,
              font_size: Optional[int] = None,
              bold_header: bool = True,
              alignment: Optional[str] = None) -> str:
    """向 Word 文档添加表格（Table Grid 样式）。

    Args:
        file_path: 文档路径
        data: 二维数据数组，如 [['A1','B1'], ['A2','B2']]
        headers: 表头列表，如 ['列1', '列2']
        font_name: 表格字体
        font_size: 字号（磅）
        bold_header: 表头是否加粗
        alignment: 表格对齐 (left/center/right)

    Returns:
        操作结果
    """
    if not data or not data[0]:
        return "[ERROR] 表格数据为空。"

    try:
        doc = _safe_open_doc(file_path)

        num_rows = len(data) + (1 if headers else 0)
        num_cols = len(data[0])

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        if alignment:
            table_align = TABLE_ALIGN_MAP.get(alignment.lower())
            if table_align:
                table.alignment = table_align

        # 写表头
        if headers:
            for j, header_text in enumerate(headers):
                if j < num_cols:
                    cell = table.rows[0].cells[j]
                    cell.text = str(header_text)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            _set_run_font(run, font_name=font_name,
                                          font_size=font_size, bold=bold_header)

        # 写数据
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[start_row + i].cells[j]
                    cell.text = str(cell_text)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            _set_run_font(run, font_name=font_name, font_size=font_size)

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 表格 ({num_rows}x{num_cols}) 已添加至 '{file_path}'"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("添加表格失败")
        return f"[ERROR] 添加表格失败: {e}"


def add_page_break(file_path: str) -> str:
    """向 Word 文档末尾添加分页符。

    Args:
        file_path: 文档路径

    Returns:
        操作结果
    """
    try:
        doc = _safe_open_doc(file_path)
        doc.add_page_break()
        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 分页符已添加至 '{file_path}'"
        return "[ERROR] 保存文档失败"
    except Exception as e:
        logger.exception("添加分页符失败")
        return f"[ERROR] 添加分页符失败: {e}"


# ────────────────────────── 格式调整接口 ──────────────────────────

def set_paragraph_format(file_path: str,
                         paragraph_index: Optional[int] = None,
                         alignment: Optional[str] = None,
                         line_spacing: Optional[float] = None,
                         space_before: Optional[int] = None,
                         space_after: Optional[int] = None,
                         first_line_indent: Optional[int] = None,
                         left_indent: Optional[int] = None,
                         right_indent: Optional[int] = None) -> str:
    """调整段落排版格式。

    Args:
        file_path: 文档路径
        paragraph_index: 段落索引（从 1 开始），不传则作用于全部段落
        其他参数同 add_paragraph 的格式参数

    Returns:
        操作结果
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        paragraphs = doc.paragraphs

        if paragraph_index is not None:
            if not 1 <= paragraph_index <= len(paragraphs):
                return f"[ERROR] 段落索引 {paragraph_index} 超出范围 (1-{len(paragraphs)})"
            target = [paragraphs[paragraph_index - 1]]
            desc = f"第 {paragraph_index} 段"
        else:
            target = paragraphs
            desc = "所有段落"

        for para in target:
            _set_para_format(para, alignment=alignment, line_spacing=line_spacing,
                             space_before=space_before, space_after=space_after,
                             first_line_indent=first_line_indent,
                             left_indent=left_indent, right_indent=right_indent)

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 已调整 '{file_path}' 中 {desc} 的格式"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("调整段落格式失败")
        return f"[ERROR] 调整段落格式失败: {e}"


def set_run_format(file_path: str, paragraph_index: int,
                   font_name: Optional[str] = None,
                   font_size: Optional[int] = None,
                   bold: Optional[bool] = None,
                   italic: Optional[bool] = None,
                   underline: Optional[bool] = None,
                   color: Optional[str] = None) -> str:
    """调整指定段落的字体格式。

    Args:
        file_path: 文档路径
        paragraph_index: 段落索引（从 1 开始）
        其他参数同 add_paragraph 的字体参数

    Returns:
        操作结果
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        if not 1 <= paragraph_index <= len(doc.paragraphs):
            return f"[ERROR] 段落索引 {paragraph_index} 超出范围 (1-{len(doc.paragraphs)})"

        para = doc.paragraphs[paragraph_index - 1]
        for run in para.runs:
            _set_run_font(run, font_name=font_name, font_size=font_size,
                          bold=bold, italic=italic, underline=underline, color=color)

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 已调整 '{file_path}' 第 {paragraph_index} 段的字体"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("调整字体格式失败")
        return f"[ERROR] 调整字体格式失败: {e}"


def set_page_margins(file_path: str,
                     top: Optional[float] = None,
                     bottom: Optional[float] = None,
                     left: Optional[float] = None,
                     right: Optional[float] = None) -> str:
    """设置 Word 文档的页面边距。

    Args:
        file_path: 文档路径
        top: 上边距（英寸）
        bottom: 下边距（英寸）
        left: 左边距（英寸）
        right: 右边距（英寸）

    Returns:
        操作结果
    """
    if not os.path.exists(file_path):
        return f"[ERROR] 文件 '{file_path}' 不存在。"

    try:
        doc = Document(file_path)
        section = doc.sections[0]

        if top is not None:
            section.top_margin = Inches(top)
        if bottom is not None:
            section.bottom_margin = Inches(bottom)
        if left is not None:
            section.left_margin = Inches(left)
        if right is not None:
            section.right_margin = Inches(right)

        if _safe_save_doc(doc, file_path):
            return f"[SUCCESS] 已设置 '{file_path}' 页面边距 (上:{top} 下:{bottom} 左:{left} 右:{right})"
        return "[ERROR] 保存文档失败"

    except Exception as e:
        logger.exception("设置页边距失败")
        return f"[ERROR] 设置页面边距失败: {e}"
